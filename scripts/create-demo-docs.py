"""
Generate 4 demo documents for fynqAI Phase 9 demo.
Run: python3 scripts/create-demo-docs.py
Output: public/demo-docs/
"""

import os
from pathlib import Path

try:
    from fpdf import FPDF
except ImportError:
    raise SystemExit("fpdf2 not installed. Run: pip install fpdf2")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx")


OUTPUT_DIR = Path(__file__).parent.parent / "fynqai" / "public" / "demo-docs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# PDF helper
# ─────────────────────────────────────────────────────────────────────────────
def make_pdf(filename: str, sections: list[dict]) -> None:
    """sections = [{"heading": str, "body": str}, ...]"""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(20, 20, 20)
    pdf.cell(0, 12, sections[0]["heading"], new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
    pdf.ln(6)

    for section in sections[1:]:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 8, section["heading"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(60, 60, 60)
        pdf.multi_cell(0, 6, section["body"])
        pdf.ln(6)

    out_path = OUTPUT_DIR / filename
    pdf.output(str(out_path))
    print(f"  ✓ {out_path.name}")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX helper
# ─────────────────────────────────────────────────────────────────────────────
def make_docx(filename: str, sections: list[dict]) -> None:
    doc = Document()

    # Title
    title = doc.add_heading(sections[0]["heading"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section in sections[1:]:
        doc.add_heading(section["heading"], level=2)
        p = doc.add_paragraph(section["body"])
        p.paragraph_format.space_after = Pt(8)

    out_path = OUTPUT_DIR / filename
    doc.save(str(out_path))
    print(f"  ✓ {out_path.name}")


# ─────────────────────────────────────────────────────────────────────────────
# File 1: supplier-agreement.pdf
# ─────────────────────────────────────────────────────────────────────────────
make_pdf("supplier-agreement.pdf", [
    {"heading": "Supplier Agreement - Version 1.2 (2022)", "body": ""},
    {"heading": "Overview",
     "body": (
         "This agreement governs the terms of supply and delivery between FynqCorp and its "
         "listed domestic and international suppliers. All parties are bound by the clauses "
         "stated herein unless superseded by a more recent written agreement."
     )},
    {"heading": "Delivery Terms",
     "body": (
         "All domestic orders have a minimum order weight of 10kg. Orders below this threshold "
         "will incur a surcharge of $15 per shipment. Bulk international orders are subject to "
         "separate freight negotiation and must be pre-approved by the logistics team."
     )},
    {"heading": "Payment Terms",
     "body": (
         "Standard payment terms are net-30 from the date of invoice. Early payment discounts "
         "of 2% are available for settlements made within 10 business days."
     )},
    {"heading": "Quality Standards",
     "body": (
         "All supplied goods must conform to ISO 9001:2015 standards. Non-conforming deliveries "
         "will be returned at the supplier's expense within 5 business days of receipt."
     )},
])

# ─────────────────────────────────────────────────────────────────────────────
# File 2: operations-manual.docx
# ─────────────────────────────────────────────────────────────────────────────
make_docx("operations-manual.docx", [
    {"heading": "Operations Manual — Revision 4 (2024)", "body": ""},
    {"heading": "Scope",
     "body": (
         "This manual applies to all domestic and regional operations managed by the FynqCorp "
         "logistics and fulfilment divisions. It supersedes all previous operational guidelines "
         "issued prior to the Q1 2024 revision."
     )},
    {"heading": "Order Requirements",
     "body": (
         "Bulk domestic orders require a minimum order weight of 50kg per shipment. "
         "This applies to all standard delivery routes. Orders that cannot meet this threshold "
         "must be consolidated with other pending orders before dispatch authorisation is granted."
     )},
    {"heading": "Warehouse Operations",
     "body": (
         "All inbound goods must be scanned within 2 hours of arrival at any FynqCorp facility. "
         "Rejected or returned stock must be held in the quarantine zone pending quality review."
     )},
    {"heading": "Fleet Management",
     "body": (
         "Domestic delivery fleets operate Monday to Saturday, 06:00-20:00 local time. "
         "Route optimisation software must be used for all multi-stop journeys exceeding 3 drops."
     )},
])

# ─────────────────────────────────────────────────────────────────────────────
# File 3: refund-policy.pdf
# ─────────────────────────────────────────────────────────────────────────────
make_pdf("refund-policy.pdf", [
    {"heading": "Refund & Returns Policy - Published March 2023", "body": ""},
    {"heading": "Policy Statement",
     "body": (
         "FynqCorp is committed to customer satisfaction. This policy outlines the procedures "
         "for product returns, refund eligibility, and timelines applicable to all consumer "
         "and business-to-business transactions."
     )},
    {"heading": "Returns & Refunds",
     "body": (
         "Customers may request a full refund within 30 days of confirmed delivery. Requests "
         "must be submitted via the customer portal at portal.fynqcorp.com. Refunds will be "
         "processed within 5-7 working days to the original payment method."
     )},
    {"heading": "Condition Requirements",
     "body": (
         "All returned goods must be in original, unopened condition unless the return is "
         "due to a fault or defect. Damaged or used items may be subject to a restocking fee "
         "of up to 20% of the purchase price."
     )},
    {"heading": "Exclusions",
     "body": (
         "Custom-manufactured items, perishable goods, and digital products are not eligible "
         "for return unless a manufacturing defect is confirmed by FynqCorp Quality Control."
     )},
])

# ─────────────────────────────────────────────────────────────────────────────
# File 4: support-guidelines.docx
# ─────────────────────────────────────────────────────────────────────────────
make_docx("support-guidelines.docx", [
    {"heading": "Customer Support Guidelines - Internal Use Only (2024)", "body": ""},
    {"heading": "Purpose",
     "body": (
         "These guidelines define how support agents should handle common customer requests "
         "including returns, refunds, complaints, and escalations. Agents are expected to "
         "follow these procedures precisely to ensure consistent service quality."
     )},
    {"heading": "Support Process",
     "body": (
         "Refund requests are accepted within 14 days of delivery only. After 14 days, "
         "store credit may be offered at management discretion. Agents should not approve "
         "cash refunds beyond this window without written authorisation from a team lead."
     )},
    {"heading": "Escalation Path",
     "body": (
         "Tier 1 support handles standard queries. Complaints unresolved after 2 contacts "
         "must be escalated to Tier 2. All escalations require a completed case summary "
         "submitted through the internal CRM before handoff."
     )},
    {"heading": "Agent Notes",
     "body": (
         "Always verify order status in the logistics portal before processing any refund or "
         "return request. Document every customer interaction in the CRM regardless of outcome."
     )},
])

print()
print(f"All 4 demo documents saved to: {OUTPUT_DIR}")
print()
print("Upload order (via app UI at /documents):")
print("  1. supplier-agreement.pdf")
print("  2. operations-manual.docx")
print("  3. refund-policy.pdf")
print("  4. support-guidelines.docx")
print()
print("Expected contradictions after upload:")
print("  • Minimum order weight: 10kg (supplier agreement) vs 50kg (operations manual)")
print("  • Refund window: 30 days (refund policy) vs 14 days (support guidelines)")
